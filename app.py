from tkinter import ttk as ttk1
import tkinter as tk
import sqlite3
from tkinter import messagebox
import os

import json
from time import gmtime, strftime
import time
ver = "1.0"
from tkinter import *

import ttkbootstrap as tb 







from decimal import *
# import pandas lib as pd


# Python3 code to select
import csv


def run_full_soft():
     

    from datetime import datetime
    global cur
    global connection
    connection = sqlite3.connect("stock.db")
    cur = connection.cursor()
    
    cur.execute(""" CREATE TABLE IF NOT EXISTS stock (

        id_code text,
        name text,
        brand text,
        qun_in_strips text,
        qun_in_one_strip text,
        ex_date text,
        batch_code text,
        price text,
        total_tabs text,
        gst text,
        cost_per_tab text



    )""")

    cur.execute("SELECT * from stock")
    fr = cur.fetchall()



    for gt in fr:
        print(gt)

    # datetime object containing current date and time
    now = datetime.now()







    def on_closing():
        import sys
        if messagebox.askyesnocancel("EXIT", "Do you want to EXIT?                                        WARNING:EXITING MAY DELETE UNSAVED DATA"):
            dt_string = now.strftime("%d/%m/%Y %H:%M:%S")
            print("date and time =", dt_string)

            os.chdir("logs")
            with open("CLOSING_LOG.log" , 'a')as f:
                f.write(   "CLOSED ON : "+dt_string + "\n") 
            RUNNING = False
            os.chdir("..")
            root.destroy()
            sys.exit()

            










    print(os.getcwd())

    os.chdir("logs")

    print(os.getcwd())

    os.chdir("..")

    #RUNCOMMANDS
    import random

    def run_stock_manager():



    




            


        def gen_id():
            
            os.chdir("codes")
            global read_id
            read_id = open("id.txt" , 'r').read()
            os.chdir("..")
            global loll
            loll = int(read_id)
            loll +=1
            
            os.chdir("codes")
            with open("id.txt" , 'w')as f:
                f.write(str(loll))
            os.chdir("..")
            
            
            








        def qr_gen(qr_codE_number):
                
            # Import QRCode from pyqrcode
                import pyqrcode
                
                from pyqrcode import QRCode


                # String which represents the QR code
                s = str(qr_codE_number)

                # Generate QR code
                url = pyqrcode.create(s)

                # Create and save the svg file naming "myqr.svg"

                os.chdir("qr_codes")
                # Create and save the png file naming "myqr.png"
                url.png(item_name_stock_add + ".png", scale = 3)

                os.chdir("..")







        def open_add_to_stock_win():

            def add_to_base_by_csv():
                
                 
                with open("stock.csv", 'r') as file:
                    csvreader = csv.reader(file)
                    for row in csvreader:
                         import re
                         
                         sl = row[0]
                         itm_name = row[1]
                         maf = row[2]
                         packing = row[4]
                         price123 = row[5]
                         batch = row[6]
                         pk = re.sub("\D", "",packing)
                         print(type(pk))
                         print(pk)
                         
                         
                         
                         ex = row[7]
                         QtY = row[8]
                         print(type(QtY))
                         print(QtY)
                         
                         gs1 = row[10]

                         fgst = re.sub("\D"  ,"",gs1)

                         print(fgst)



                         
                         print(price123)
                         print(type(price123))
                         total_tabs_i = (Decimal(pk) * Decimal(QtY))

                         cost_per_t = (Decimal(price123) / Decimal(pk))

                         roundedNumber = round(cost_per_t, 2)

                         print(roundedNumber)





                         cur.execute("INSERT INTO stock VALUES (? , ? , ? , ? , ? , ? , ?, ? , ? , ? , ?)" , (sl, itm_name , maf ,QtY , pk , ex , batch , str(price123) , str(total_tabs_i) , str(fgst) , str(roundedNumber)))
                         connection.commit()
                    messagebox.showinfo("" , "SUCESSFUL ADDING")















            def add_to_data_base():
                gen_id()
                global item_name_stock_add
                
                code_stock_add = str(loll)
                int(code_stock_add)
                item_name_stock_add = entry1_root2_sub.get()
                item_qun_stock_add = sp.get()
                itm_ex_date = entry3_root2_sub.get()
                itm_pricE_stock_add = entry2_root2_sub.get()
                batch = entry4_root2_sub.get()
                qun_in_one_stp = entry5_root2_sub.get()
                
                brand = entry8_root2_sub.get()

                total_tabs = (int(item_qun_stock_add) * int(qun_in_one_stp))

                gst = entry6_root2_sub.get()

                cost_per_tab = (Decimal(itm_pricE_stock_add) / Decimal(qun_in_one_stp))


                cur.execute("INSERT INTO stock VALUES (? , ? , ? , ? , ? , ? , ?, ? , ? , ? , ?)" , (code_stock_add, item_name_stock_add , brand ,item_qun_stock_add , qun_in_one_stp , itm_ex_date , batch , str(itm_pricE_stock_add) , str(total_tabs) , str(gst) , str(cost_per_tab)))

                connection.commit()

                cur.execute("SELECT * FROM stock")
                print(cur.fetchall())


                qr_gen(qr_codE_number=str(item_name_stock_add))

                messagebox.showinfo('' , 'SAVED SUCESSFULLY')

                root2_sub.destroy
                root2.destroy



            
                
                




            


            global root2_sub
            root2_sub = tb.Toplevel()
            root2_sub.geometry("1200x600")
            root2_sub.title("ADD")

            label1_root2_sub = tk.Label(   root2_sub,text="Item Name")
            label1_root2_sub.place(x=1 , y=1)

            entry1_root2_sub = tk.Entry(root2_sub)
            entry1_root2_sub.place(x=70 , y=1)


            label2_root2_sub = tk.Label(   root2_sub,text="PRICE")
            label2_root2_sub.place(x=1 , y=50)

            entry2_root2_sub = tk.Entry(root2_sub)
            entry2_root2_sub.place(x=100 , y=50)

            label9_root2_sub = tk.Label(   root2_sub,text="GST in %")
            label9_root2_sub.place(x=1 , y=100)

            entry6_root2_sub = tk.Entry(root2_sub)
            entry6_root2_sub.place(x=100 , y=100)







            label3_root2_sub = tk.Label(   root2_sub,text="QUANTITY of strips/packs(for liquids or others)")
            label3_root2_sub.place(x=200 , y=1)


            
            sp = tk.Spinbox(root2_sub, from_= 0, to = 999999999)
            sp.place(x=200 , y=25)       


            label3_root2_sub = tk.Label(   root2_sub,text="BATCH CODE / LOT CODE : ")
            label3_root2_sub.place(x=500 , y=1)

            entry4_root2_sub = tk.Entry(root2_sub)
            entry4_root2_sub.place(x=500 , y=50)

            label4_root2_sub = tk.Label(   root2_sub,text="TABS / CAPS / (ml for liquids)in one pack/strip : ")
            label4_root2_sub.place(x=750 , y=1)

            entry5_root2_sub = tk.Entry(root2_sub)
            entry5_root2_sub.place(x=750 , y=50)




            label2_root2_sub = tk.Label(   root2_sub,text="EXP DATE:")
            label2_root2_sub.place(x=500 , y=100)

            entry3_root2_sub = tk.Entry(root2_sub)
            entry3_root2_sub.place(x=500 , y=117)

            label8_root2_sub = tk.Label(   root2_sub,text="COMPANY:")
            label8_root2_sub.place(x=500 , y=200)

            entry8_root2_sub = tk.Entry(root2_sub)
            entry8_root2_sub.place(x=500 , y=250)





            def open_barcode():
                os.chdir("qr_codes")
                import webbrowser
                webbrowser.open(str(item_name_stock_add)+ ".png")
                os.chdir("..")





            btn1_root2_sub = tk.Button(root2_sub , text="SUBMIT" , command=add_to_data_base)
            btn1_root2_sub.place(x=1 , y=80)

            
            btn2_root2_sub = tk.Button(root2_sub , text="OPEN_QRCODE" , command=open_barcode)
            btn2_root2_sub.place(x=100 , y=80)

            btn3_root2_sub = tk.Button(root2_sub , text="ADD VIA CSV" , command=add_to_base_by_csv)
            btn3_root2_sub.place(x=200 , y=80)



            root2_sub.mainloop()

            


        def view_stock():
          
          def get_data():
               

                 
               for i567 in list_stock.curselection():
                        full_nam_data= list_stock.get(i567)

               cur.execute("SELECT DISTINCT * FROM stock WHERE name IN ( ? )", [full_nam_data])
               ot1 = cur.fetchall()
               
               


                

               root2_sub_3 = tb.Window(themename="solar")
               root2_sub_3.title("VIEW STOCK_DATA")
               root2_sub_3.geometry("1010x500")

               tx_bx = tk.Text(root2_sub_3 , height=15 , width=210)
               tx_bx.pack()
               tx_bx.insert(END , "ID , ITEM NAME , MANUFTURE, QUNTITY IN STRIPS, QUNTITY IN ONE STRIP, EXPIRY DATE, BATCH , PRICE , TOTAL TABS , GST" + "\n")
               for i124665 in ot1:
                    
                tx_bx.insert(END , i124665)







               root2_sub_3.mainloop()

                
                        
               
          





          root2_sub_2 = tb.Toplevel()
          root2_sub_2.title("VIEW STOCK")
          root2_sub_2.geometry("1500x500")
          st_ls_ct = 0

          label1_root_sub_2 = tk.Label(root2_sub_2 , text="")
          label1_root_sub_2.pack(pady=8)

          list_stock = tk.Listbox(root2_sub_2 , height=20 , width=200)
          list_stock.pack()
          cur.execute("SELECT * FROM stock")
          ot = cur.fetchall()

          for nm_st in ot:

                st_ls_ct += 1
                list_stock.insert(st_ls_ct , nm_st[1])



          root2_sub_2_btn = tk.Button(root2_sub_2 , text="GET DATA" , command=get_data).pack()

                



          root2_sub_2.mainloop()























            




        global root2

        root2 = tb.Toplevel()
        root2.title("T-PHARMA | STOCK MANAGER")
        root2.geometry("800x500")



        btn1_root2 = tk.Button(root2 , text="ADD TO STOCK" , command=open_add_to_stock_win)
        btn1_root2.pack()

        btn1_root3 = tk.Button(root2 , text="VIEW ALL STOCK" , command=view_stock)
        btn1_root3.pack()




        root2.mainloop()





















    def run_bill_manager():
        global options
        options = []
        options.append("SELECT")


        

        


        
        global lb_ct
        lb_ct = 0
        root3 = tb.Toplevel()
        root3.title("T-PHARMA | BILL MANAGER")
        root3.geometry("800x500")






        def open_bill_maker():
            from ttkwidgets.autocomplete import AutocompleteEntry
            root3_sub = tb.Toplevel()


            label3_root3_sub = tk.Label(root3_sub , text="PRICE EXCLUDING GST = ")
            label3_root3_sub.place(x=1 , y=10)

            label99_root3_sub = tk.Label(root3_sub , text="PRICE INCLUDING GST = ")
            label99_root3_sub.place(x=1 , y=30)

            label991_root3_sub = tk.Label(root3_sub , text="GST = ")
            label991_root3_sub.place(x=1 , y=50)








            lb = Listbox(root3_sub)




            global listofitm
            listofitm = []
            cur.execute("SELECT * FROM stock")
            

            output = cur.fetchall()

            for l in output:
                global lb_ct
                lb_ct +=1

                listofitm.append(l[1])
                


            print(listofitm)









        # Function for checking the
    # key pressed and updating
    # the listbox
            t = []

            entry = AutocompleteEntry(
            root3_sub, 
            width=30, 
            font=('Times', 18),
            completevalues=t
            )
            
            entry.place(x=500 , y=30)


            def checkkey(event):
                print("RUNNED CHECK KEY")
       
                value = event.widget.get()
                print(value)
                
                # get data from l
                if value == '':
                    data = l
                else:
                    data = []
                    for item in l:
                        if value.lower() in item.lower():
                            data.append(item)                
            
                # update data in listbox
                update(data)
            
            
            def update(data):
                
                # clear previous data
                lb.delete(0, 'end')
            
                # put new data
                for item in data:
                    lb.insert('end', item)
            
            
            # Driver code
            l = listofitm
            

            entry.bind('<Key>', checkkey)

                        
































                
                #creating text box 
            lb = Listbox(root3_sub , width=40 , font=("IMPACT"))
            lb.place(x=500 , y=70)


                    

            global slcout
            global lcount
            global total_MRP
            global total_GST_AMM
            global options
            lcount = 0
            slcout = 0
            total_MRP = 0
            global ls_co
            ls_co = 0
            total_GST_AMM = 0
            global data_to_add
            data_to_add = []
            


            def get_price():
                    print("RUNNED GET PRICE")

                    
                    for a in lb.curselection():
                        itm_name = lb.get(a)
                    
                    
                    cur.execute("SELECT DISTINCT * FROM stock WHERE name IN ( ? )", [itm_name])

                    ckpr = cur.fetchall()


                    for bb in ckpr:
                         global ls_co
                         

                         ls_co +=1 
                         
                         

  
                         ls.insert(ls_co , bb[6])
                         ls2.insert(ls_co , bb[7])

                         
                    

                         

                         











                    for ii in ckpr:
                                
                                gst_per_cent_num = ii[9]

                                pris = ii[7]


                                def gi():
                                     global GST
                                     GST =    Decimal(pris) + (Decimal(pris)*(Decimal(gst_per_cent_num)/100))

                                if gst_per_cent_num == "":
                                     GST = pris
                                     gst_per_cent_num = "0"
                                else:
                                     gi()
                                    


                                
                                




                                
                                
                                
                                

                                
            

                                label3_root3_sub.config(text= "COST OF ONE STRIP GST EXCLUDING :" + ii[7])
                                label99_root3_sub.config(text="COST OF ONE STRIP INCLUDING GST: " + str(GST))
                                label991_root3_sub.config(text="GST = " + str(gst_per_cent_num))
                                

            def run_get(event):
                 print("RUNNED RUN GET")
                 ls.delete(0 ,END)
                 ls2.delete(0 , END)


                 
                     


                 


                 get_price()
                 

            
            entry.bind('<Key>', checkkey)







 
                        
             



            lb.bind('<<ListboxSelect>>', run_get)



            def ins():
                    
                    
                        






                    for a in lb.curselection():
                        itm_name_entered = lb.get(a)

                    
                    options.clear()

                    
                    
                    global slcout
                    global total
                    

                    qun_by_tab_entered = entry4_root3_sub.get()

                    if qun_by_tab_entered == "":
                        qun_by_tab_entered = "1"

                    

                    cur.execute("SELECT DISTINCT * FROM stock WHERE name IN ( ? )", [itm_name_entered])
                    output = cur.fetchall()
                    print(output)

                    for i in ls.curselection():
                        bth = ls.get(i)
                        print(bth)

                    
                    for i1 in ls2.curselection():
                        pr123456 = ls2.get(i1)
                        






                    for getting in output:
                         pri = getting[7]
                         qins = getting[4]

                         
                         
                         brand = getting[2]
                         cost_per_tablet = getting[10]
                         gst_get_cent = getting[9]


                         cost_per_tab5555 = (Decimal(pr123456) / Decimal(qins))
                         print(cost_per_tab5555)

                         cpt_r = round(int(cost_per_tab5555) , 2)

                         print(cpt_r)

                         if cpt_r == 1:
                             cpt_r = pr123456
                         



                         





























                         

                         
                         total_tabs_get = getting[8]
                         ex_date_get = getting[5]
                         

                         print(brand)
                         
                         print(total_tabs_get)
                         print(ex_date_get)
                         print(qun_by_tab_entered)
                         print(cpt_r)
                         
                         
                         fin_pri = (Decimal(qun_by_tab_entered) * Decimal(cpt_r))

                         
                         

                         

                         

                         
                         
                         def gg():
                              global total_GST_AMM
                              total_GST_AMM += (int(fin_pri)*(int(gst_get_cent)/100))
                              

                              
                              
                              
                              

                         if gst_get_cent == "":
                              gst_get_cent = "0"
                              global total_GST_AMM
                              
                              total_GST_AMM += fin_pri

                         else:
                              
                              gg()
                         
                         
                         

                         
                         
                         global total_MRP
                         
                         total_MRP += int(fin_pri)

                         update_total_tab = (int(total_tabs_get) - int(qun_by_tab_entered))
                         

                         

                         print(data_to_add)
                    
                    slcout += 1
                    data_to_add.append([  str(slcout) , itm_name_entered , brand ,  "₹"+str(fin_pri) , str(qun_by_tab_entered) , str(bth) , ex_date_get , str(gst_get_cent)])
                    sh_ls.insert(END , "ITEM:" + itm_name_entered + '    |    ' + "MAF.:" + str(brand) + '    |    ' + "MRP:" + str(fin_pri) + '    |    ' + "QUANTITY:" + qun_by_tab_entered + '    |    ' + "BATCH:" + bth + '    |    ' + "EX_DATE:" + ex_date_get + '    |    ' + "GST%:" + str(gst_get_cent))
                    
                    



            def save():
                 global drop
                 global pam
                 pam = menu.get()
                 from docx.shared import Pt
                 os.chdir("settings")
                 dag = open("dag" , "r").read()
                 D_l = open("DL" , "r").read()
                 GsTiN = open("GSTIN" , "r").read()
                 os.chdir("..")
                     
                 

                 
                 cust_name = entry5_root3_sub.get()
                 from datetime import datetime

                 now = datetime.now() # current date and time

                 



                 global date_time1
    
                 date_time = now.strftime("%d/%m/%Y")
                 date_time1 = now.strftime("%d.%m.%Y")

                 print(str(date_time1))

                 os.chdir("bills")
                 if not os.path.exists(str(date_time1)):
                  
                  os.makedirs(str(date_time1))

                 os.chdir("..")

                 








                 from docx import Document

                 doc = Document()

                 from docx.shared import Inches, Cm

                 section = doc.sections[0]
                 section.left_margin = Cm(1.0)
                 section.right_margin = Cm(1.0)
                 section.top_margin = Cm(1.0)

                 os.chdir("settings")
                 nmpr = open("pharmacy_name.txt" , 'r').read()
                 # Import docx NOT python-docx
                 os.chdir("..")
                 doc.add_heading(nmpr, 3)
                 para = doc.add_heading("DL NO: " + D_l+ "  GSTIN: " + GsTiN+ "  DAG NO:   "  + dag, 4)


                 para = doc.add_heading("INVOICE")

                 para.paragraph_format.space_before = Pt(1)
                 para.paragraph_format.space_after = Pt(0.2)





                 


                 top0 = [["Date:" + date_time , "" ] ,
                         ["Customer Name:" + cust_name , ""],
                         
                         
                         
                         
                         ]

    

                 table1 = doc.add_table(rows=1 , cols=2)


                 

                 for one , two  in top0:
                    cells1 = table1.add_row().cells
                    cells1[0].text = one
                    cells1[1].text = two



                

                 table_header = ["SL.NO" , "Item" , "Maf." , "PRICE" , "Quantity" , "Batch" , 'EXP DATE' , "GST%"]

                 table = doc.add_table(rows=2 , cols=8)


                 for i in range(8):
                    table.rows[0].cells[i].text = table_header[i]


                 for sl_num , item1245 , company , price , quntity , batch , exp_date124 , gst214 in data_to_add:
                    cells = table.add_row().cells
                    cells[0].text = sl_num
                    cells[1].text = item1245
                    cells[2].text = company
                    cells[3].text = price
                    cells[4].text = quntity
                    cells[5].text = batch
                    cells[6].text = exp_date124
                    cells[7].text = gst214


                 for row in table.rows:
                    for cell in row.cells:

                        paragraphs = cell.paragraphs

                        for paragraph in paragraphs:
                          
                          for run in paragraph.runs:

                            font = run.font
                            font.size= Pt(8)




                 table.style = 'Colorful List'

                 global total_GST_AMM
                 global total_MRP



                 
                 if total_MRP == total_GST_AMM:
                      net_pay_able = total_MRP
                      total_GST_AMM = 0

                 else:
                      r_g = round(total_GST_AMM)
                      net_pay_able = (Decimal(r_g) + Decimal(total_MRP))

                
                 r_net_pay_able = round(net_pay_able , 2)
                 
                 


                 global usr_logged_in
                 
                 
                 bott = [[ "Total MRP: ₹" + str(total_MRP), "GST AMMOUNT: ₹" + str(total_GST_AMM) ,"Ammount Payable: ₹" +str(r_net_pay_able)],

                         ["Billed By:" + usr_logged_in ,"Payment: " + pam ,""],

                         ["Note: Price will be as per Quantity NOT MRP." , "Quantity is in loose NOT pack" , ""]
                         
                         
                         
                         ]

    

                 table2 = doc.add_table(rows=1 , cols=3)
                

                 

                 for one , two , three in bott:
                    cells1 = table2.add_row().cells
                    cells1[0].text = one
                    cells1[1].text = two
                    cells1[2].text = three


                 for row1 in table2.rows:
                    for cell1 in row1.cells:

                        paragraphs1 = cell1.paragraphs

                        for paragraph1 in paragraphs1:
                          
                          for run1 in paragraph1.runs:

                            font1 = run1.font
                            font1.size= Pt(8)



















                    

                    

                 
                 os.chdir("bills")

                 
                 entry5_root3_sub.delete(END)
                 entry4_root3_sub.delete(END)






                 os.chdir(str(date_time1))
 
# Now save the document to a location
                 doc.save(cust_name + ".docx")
                 import webbrowser
                 
                 import print1
                 webbrowser.open(cust_name + ".docx")
                 #print.print_file(fl_nm=cust_name+".docx")
                 

                 os.chdir("..")
                 os.chdir("..")



                 



    # Create an instance of a word document
 
                                        

                                    
                                    
                            






                            




                            

                        
                    








                            
                    


   





            
            root3_sub.geometry("1500x700")
            root3_sub.title("CREATE BILL ---- WIZARD ------ ")

            label2_root3_sub = tk.Label(root3_sub , text="SEARCH HERE")
            label2_root3_sub.place(x=500 , y=1)







            label4_root3_sub = tk.Label(root3_sub , text="QUANTITY")
            label4_root3_sub.place(x=730 , y=1)

            entry4_root3_sub = tk.Entry(root3_sub)
            entry4_root3_sub.place(x=800 , y=1)

            btn1_root_3_sub = tb.Button(root3_sub , text="INSERT" , command=ins ,bootstyle="danger-outlined")
            btn1_root_3_sub.place(x=900 , y=1)




            
            btn3_root_3_sub = tb.Button(root3_sub , text="SAVE & PRINT(WITH DEFAULT PRINTER [TO SET DEFAULT PRINTER GO TO PC SETTINGS])" , command=save,bootstyle="danger-outlined")
            btn3_root_3_sub.place(x=900 , y=580)




            
            label5_root3_sub = tk.Label(root3_sub , text="Customer Name:")
            label5_root3_sub.place(x=900 , y=500)

            entry5_root3_sub = tb.Entry(root3_sub)
            entry5_root3_sub.place(x=1000 , y=500)





            
            # datatype of menu text

            # Create Dropdown menu
            menu= StringVar()
            menu.set("---SELECT PAYMENT---")

            #Create a dropdown Menu
            drop= OptionMenu(root3_sub, menu,"CASH" , "CARD" , "UPI" , "NET BANKING")
            drop.place(x=950 , y=470)
                        
            # Create button, it will change label text

            
            # Create Label






            # Dropdown menu options


            
            # initial menu text

            
            global ls
                 
            ls = Listbox(root3_sub , exportselection=False)
            ls.place(x=250 , y=1)

            ls2 = Listbox(root3_sub , exportselection=False)
            ls2.place(x=370 , y=1)

        
    
            
            # Create Dropdown menu

            sh_ls = tk.Listbox(root3_sub , width=150)
            sh_ls.place(x=1 ,y=400)








    







            def run_ins(event):
                 ins()



















            root3_sub.bind('<Return>', run_ins)

            root3_sub.mainloop()









        btn1_root_3 = tb.Button(root3 , text="Create Bill" , command=open_bill_maker  ,bootstyle="danger-outlined")
        btn1_root_3.pack(pady=50)






        root3.mainloop()

        















    def run_settings():
        root4 = tk.Toplevel()
        root4.title("T-PHARMA | SETTINGS")
        root4.geometry("800x500")





        



        root4.mainloop()









    #RUNCOMMANDS ENDHERE


    os.chdir("settings")
    phar_name = open("pharmacy_name.txt" , 'r').read()
    os.chdir("..")
    #INTINILIZAITION AND TIME
    root = tb.Toplevel()

    root.attributes('-fullscreen', False)
    root.title("T - pharma   |   TSOFT    | VER-1.0")
    root.geometry("1300x800")
    












    
    root.iconbitmap("")

    root.attributes('-fullscreen', True)



    phar_name_DISPLAY = tk.Label(root , text=phar_name , font=("Ariel" , 10))
    phar_name_DISPLAY.pack(anchor="center")



    #ENDING HERE

    from PIL import ImageTk, Image

    img = ImageTk.PhotoImage(Image.open("bag.jpg"))
  
# reading the image
    panel = tk.Label(root, image = img)
    
    # setting the application
    panel.pack(side = "bottom", fill = "both",
           expand = "yes")
    
    # Create a window




# Set the resolution of window


# Allow Window to be resizable




    






    def logg_ouut():
         if messagebox.askyesnocancel("LOGOUT" , "ARE YOU SURE TO LOGOUT ?") == True:
              import subprocess
              import sys
              subprocess.call([r'logout.bat'])

              sys.exit()
              


    #MAIN CODE STARTING:

    label1 = tk.Label(root , text="T-PHARMA  |   " + "VERSION:" + ver , font=("Terminal" , 15))
    label1.place(x=1 , y=1)

    btn1 = tb.Button(root , text="OPEN STOCK MANAGER" , bootstyle="warning" , command=run_stock_manager )
    btn1.place(x=10 , y=50)

    btn2 = tb.Button(root , text="OPEN BILL MANAGER" , bootstyle="warning",command=run_bill_manager)
    btn2.place(x=10 , y=80)

         

    

            
         

    btn3 = tb.Button(root , text="SETTINGS" , command=run_settings,bootstyle="warning")
    btn3.place(x=10 , y=110)

    
    



    btn4 = tb.Button(root , text="OPEN SALES MANAGER",bootstyle="danger-outlined")
    btn4.place(x=10 , y=140)


    btn5 = tb.Button(root , text="LOGOUT" , command=logg_ouut,bootstyle="danger-outlined")
    btn5.pack(      )
    
    btn6 = tb.Button(root , text="EXIT" , command=on_closing,bootstyle="danger-outlined")
    btn6.pack(anchor="ne")




    root.protocol("WM_DELETE_WINDOW", on_closing)

    #RUNNING WIN <
    root.mainloop()
    #RUNNING WIN >













def open_admin_window_for_user_management():
               
     connection1 = sqlite3.connect("userdata.db")
     cur1 = connection1.cursor()
     def reg_usr():
          from tkinter.simpledialog import askstring
          from tkinter.messagebox import showinfo
          name_usr = askstring('Name', 'ENTER USER NAME?')
          password_set = askstring('PASSWORD', 'ENTER PASSWORD')

          cur1.execute(""" CREATE TABLE IF NOT EXISTS user_data (user_name text , password text)""")
          cur1.execute("INSERT INTO user_data VALUES (? , ? )" , (name_usr , password_set))
                       
          
          connection1.commit()
          messagebox.showinfo("" , "SAVED")

        
     

     win = tk.Tk()
     win.geometry("800x500")
     win.title("Admin window")

     button = tk.Button(win , text="REGISTER NEW USER" , command=reg_usr)
     button.place(x=50 , y=50)

     win.mainloop()








     


def validate_for_admin():
    


    def runveri()  :

         
        import tkinter as tk

        from functools import partial
        import random
        import os
        import smtplib
        import random
        from time import strftime


        import time


        import asyncio
        import os



        sysver = "1.2"







        #genarating otp :



        alpha = "abcdefghichlmopqrstuvwxyz"

        nums = "0123456789"

        sign = "!@#$%^&*()"



        for i in range(1,2):



            otp = random.choice( nums )

            otp2 = random.choice( nums )

            otp3 = random.choice( nums )

            otp4 = random.choice( nums )

            otp5 = random.choice( nums )

            otp6 = random.choice( nums )




            main = (otp + otp2 + otp3 + otp4 + otp5 + otp6)



        #defs

































        
            


        
                











            










            
            








        def otpvalidation(otpenter):

            x = otpenter.get()

            if main == x:
                
                root.destroy()


                open_admin_window_for_user_management()
                
                





            else:
                print("OTP NOT VALID")
                otpvalidation()

            


















            










        def validateLoginf():
                os.chdir("settings")
                
                admin_email = open("admin_email" , 'r').read()
                
                mymail = admin_email

                print("Sending OTP to the provided email address : " + mymail)
                

                tkWindow.destroy()



                s = smtplib.SMTP('smtp.gmail.com', 587)

        # start TLS for security
                s.starttls()


        # Authentication
                s.login("mosaraf.hal@gmail.com", "gccjvlforyzzoqei")

        # message to be sent



        # sending the mail
                s.sendmail("mosaraf.hal@gmail.com", mymail ,  main )




                print("Sucessfully  Sent OTP| Please Check | OTP SENDER ACOUNT = mosaraf.hal@gmail.com")

        # terminating the session
                s.quit()

                os.chdir("..")





        #window












        validateLoginf()


        root = tk.Tk()

        root.geometry("1000x300")

        root.title("PLEASE ENTER OTP")

        otpenter = tk.StringVar()

        lbl1 = tk.Label(root , text="ENTER OTP :")
        lbl1.place(x=5,y=10)
        ebl1 = tk.Entry(root,textvariable=otpenter)

        
        lbl2 = tk.Label(root , text="(OTP WILL BE SENT TO THE EMAIL ADDRESS OF THE ADMIN)" , font=("Terminal" , 15))
        lbl2.place(x=400,y=10)

                

        ebl1.place(x=100 , y=10)

        otpvalidation = partial(otpvalidation , otpenter)




        sub = tk.Button(root, text="Submit", command=otpvalidation)

        sub.place(x= 100 , y= 100)

        root.mainloop()             
            


    if messagebox.askokcancel("DO YOU WANT TO CONTINUE? ? ?" , "THIS ACTION WILL SEND AN OTP TO THE ADMIN EMAIL FOR VERIFICATION AND SECURITY PURPOSES") == True:
        runveri()









from tkinter import *
from functools import partial











def main():
	# destroy splash window
    splash_root.destroy()

	# Execute tkinter







def validateLogin(username, password):
        connection1 = sqlite3.connect("userdata.db")
        cur1 = connection1.cursor()
        
        
        pasw = password.get()
        ur = username.get()



        cur1.execute("SELECT DISTINCT * FROM user_data WHERE user_name IN ( ? )", [ur])
        get1 = cur1.fetchall()
        print(get1)

        for i135 in get1:
             
             p = i135[1]

             if p == pasw:
                






                global usr_logged_in

                usr_logged_in = ur
                

               
                
                
                tkWindow.state(newstate='iconic')
                run_full_soft()

        else:
             messagebox.showerror("" , "USERNAME OR PASSWORD IS INCORRECT")


                    
                       
                
                
             




            
#window

try:
     


# import the following modules
    import json
    import requests




        # Clear the above screen


        
        # URL from where we will fetch the data
    url = "https://uselessfacts.jsph.pl/random.json?language=en"
        
        # Use GET request
    response = requests.request("GET", url)
        
        # Load the request in json file
    data = json.loads(response.text)
                
        # we will need 'text' from the list, so
        # pass 'text' in the list
    
    useless_fact = data['text']
    print(useless_fact)
        
                
    
















except Exception:
     useless_fact = "   FAILED TO GIVE FACT BECAUSE :   NO INTERNET "
     print("GOT ERROR BUT CONTINUING")
     




























# Import module
from tkinter import *
global splash_root
# Create object
splash_root = tk.Tk()
splash_root.attributes('-fullscreen', True)
splash_root.config(bg="blue")
# Adjust size
splash_root.geometry("100x100")



# Set Label
splash_label = Label(splash_root, text="T-PHARMA", font=("Terminal" , 20))
splash_label.place(x=500 , y=100)

splash_label1 = Label(splash_root, text="Loading.......", font=("Terminal" , 20))
splash_label1.place(x=500 , y=200)

splash_label2 = Label(splash_root, text="Please Wait", font=("Terminal" , 20))
splash_label2.place(x=500 , y=300)

splash_label3 = Label(splash_root, text="TSOFT © 2023    DEVLOPER-TAMGHNA IQUEBAL   CONTACT - tamghnaiquebal@gmail.com", font=("Terminal" , 17))
splash_label3.place(x=500 , y=800)


\
splash_label4 = Label(splash_root, text="FUN FACT BY TSOFT:" + useless_fact, font=("Ariel" , 15))
splash_label4.place(x=500 , y=400)



# main window function
splash_root.after(1000, main)
















# Set Interval


# Execute tkinter
mainloop()





     



def run_valid(event):
        
    validateLogin()




tkWindow = tb.Window(themename="morph")
tkWindow.attributes('-fullscreen', True)
tkWindow.geometry('400x400')  
tkWindow.title('PLEASE LOGIN NOW!')


tkWindow.bind('<Return>', run_valid)
usernamLabel = tb.Label(tkWindow, text="----LOGIN NOW----" , font=("Terminal" , 17) ).pack(pady=20)
    #username label and text entry box
usernameLabel = tb.Label(tkWindow, text="User Name", font=("IMPACT" , 15)).pack()  
username = StringVar()
usernameEntry = tb.Entry(tkWindow, textvariable=username, font=("IMPACT" , 15)).pack()

    #password label and password entry box
passwordLabel = tb.Label(tkWindow,text="Password", font=("IMPACT" , 15)  )
passwordLabel.pack()
password = StringVar()
passwordEntry = tb.Entry(tkWindow, textvariable=password, show='*', font=("IMPACT" , 15)).pack()

validateLogin = partial(validateLogin, username, password)

    #login button
loginButton = tb.Button(tkWindow, text="Login", command=validateLogin,bootstyle="danger-outlined").pack(pady=20)               



ADMIN_button = tb.Button(tkWindow , text="ADMIN MODE" , command=validate_for_admin,bootstyle="danger-outlined")
ADMIN_button.pack(pady=10)
def exit12():
    import sys
    sys.exit()

EXIT_button = tb.Button(tkWindow , text="EXIT" , command=exit12 ,bootstyle="danger-outlined")
EXIT_button.pack(anchor="center")

tkWindow.mainloop()
  # Put the facts in the blue colour\



















